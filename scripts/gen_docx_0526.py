"""
Generate learning notes Word document for LongCut Stage 5
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Heading styles
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Arial'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.font.bold = True
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.font.bold = True
    else:
        heading_style.font.size = Pt(13)
        heading_style.font.bold = True

# ===== Title =====
doc.add_heading('LongCut 学习笔记 \u2014 阶段五：React 进阶', level=1)

p = doc.add_paragraph()
p.add_run('日期：2026-05-25 ~ 2026-05-26').font.size = Pt(10)
p = doc.add_paragraph()
p.add_run('涵盖：React Hooks\u6DF1\u5165 | \u81EA\u5B9A\u4E49Hook | Zustand | Server Components | API Routes').font.size = Pt(10)

doc.add_page_break()

# =================== Phase 5.1 ===================
doc.add_heading('\u4E00\u3001\u9636\u6BB5 5.1 \u2014 React Hooks \u6DF1\u5165', level=1)

doc.add_heading('1.1 useEffect\uff1A\u6761\u4EF6\u6267\u884C\u5668', level=2)
p = doc.add_paragraph()
p.add_run('\u6838\u5FC3\u7406\u89E3\uFF1A').bold = True
p.add_run('useEffect \u4E0D\u662F\u7F13\u5B58\uFF0C\u662F\u6761\u4EF6\u6267\u884C\u5668\u2014\u2014\u63A7\u5236\u4EE3\u7801\u201C\u4EC0\u4E48\u65F6\u5019\u6267\u884C\u201D\u3002')

doc.add_paragraph('\u4F9D\u8D56\u6570\u7EC4\u4E3A []\uFF1A\u7EC4\u4EF6\u6302\u8F7D\u540E\u6267\u884C\u4E00\u6B21')
doc.add_paragraph('\u4F9D\u8D56\u6570\u7EC4\u4E3A [dep]\uFF1A\u6BCF\u6B21 dep \u53D8\u5316\u65F6\u6267\u884C')
doc.add_paragraph('\u4E0D\u4F20\uFF1A\u6BCF\u6B21\u6E32\u67D3\u540E\u90FD\u6267\u884C\uFF08\u5371\u9669\uFF0C\u5BB9\u6613\u6B7B\u5FAA\u73AF\uFF09')

p = doc.add_paragraph()
p.add_run('\u5173\u952E\u533A\u5206\uFF1A').bold = True
p.add_run('\u201C\u6267\u884C useEffect \u91CC\u7684\u4EE3\u7801\u201D \u2260 \u201C\u7EC4\u4EF6\u91CD\u65B0\u6E32\u67D3\u201D\u3002effect \u5728\u6E32\u67D3\u5B8C\u6210\u540E\u6267\u884C\u3002')

doc.add_heading('1.2 useMemo\uff1A\u503C\u7F13\u5B58', level=2)
doc.add_paragraph('\u5B58\u50A8\u7684\u662F\u8BA1\u7B97\u7ED3\u679C\uFF08\u503C\uFF09\uFF0C\u4E0D\u662F\u51FD\u6570\u3002\u4F9D\u8D56\u4E0D\u53D8\u5219\u590D\u7528\u4E0A\u6B21\u7ED3\u679C\u3002')

doc.add_heading('1.3 useCallback\uff1A\u51FD\u6570\u7F13\u5B58', level=2)
doc.add_paragraph('\u5B58\u50A8\u7684\u662F\u51FD\u6570\u672C\u8EAB\uFF0C\u4FDD\u8BC1\u5F15\u7528\u4E0D\u53D8\uFF0C\u8BA9 === \u901A\u8FC7\u3002\u907F\u514D\u5B50\u7EC4\u4EF6\u767D\u91CD\u6E32\u67D3\u3002')

doc.add_heading('1.4 \u5F15\u7528\u76F8\u7B49\u4E0E\u91CD\u6E32\u67D3', level=2)
p = doc.add_paragraph()
p.add_run('React \u7528 === \u6BD4\u8F83 props\u3002\u56E0\u4E3A () => {} \u6BCF\u6B21\u90FD\u662F\u65B0\u5BF9\u8C61\uFF1A')
doc.add_paragraph('() => {} === () => {}  // \u59CB\u7EC8 false')
doc.add_paragraph('\u6240\u4EE5\u6BCF\u6B21\u7236\u7EC4\u4EF6\u6E32\u67D3\u2192\u4F20\u7ED9\u5B50\u7EC4\u4EF6\u7684\u51FD\u6570\u662F\u201C\u65B0\u7684\u201D\u2192\u5B50\u7EC4\u4EF6\u767D\u91CD\u6E32\u67D3\u3002useCallback \u5C31\u662F\u89E3\u51B3\u8FD9\u4E2A\u95EE\u9898\u3002')

doc.add_page_break()

# =================== Phase 5.2 ===================
doc.add_heading('\u4E8C\u3001\u9636\u6BB5 5.2 \u2014 \u81EA\u5B9A\u4E49 Hook', level=1)

doc.add_heading('2.1 \u81EA\u5B9A\u4E49 Hook \u89C4\u5219', level=2)
doc.add_paragraph('\u51FD\u6570\u540D\u5FC5\u987B\u4EE5 use \u5F00\u5934\uFF08React \u7684\u7EA6\u5B9A\uFF09')
doc.add_paragraph('\u5185\u90E8\u53EF\u4EE5\u8C03\u7528\u5176\u4ED6 Hook\uFF08useState\u3001useEffect \u7B49\uFF09')
doc.add_paragraph('Hook \u8C03\u7528\u4E0D\u80FD\u653E\u5728\u6761\u4EF6\u5206\u652F\u91CC\uFF08React \u4F9D\u8D56\u8C03\u7528\u987A\u5E8F\u8BC6\u522B\u72B6\u6001\uFF09')

doc.add_heading('2.2 useDebounce \u5B9E\u73B0\u539F\u7406', level=2)
p = doc.add_paragraph()
p.add_run('\u573A\u666F\uFF1A').bold = True
p.add_run('\u641C\u7D22\u6846\u8F93\u5165\u65F6\uFF0C\u4E0D\u5728\u6BCF\u6B21\u6309\u952E\u90FD\u53D1\u8BF7\u6C42\uFF0C\u800C\u662F\u7B49\u7528\u6237\u505C\u6B62\u8F93\u5165 300ms \u540E\u518D\u53D1\u3002')

p = doc.add_paragraph()
p.add_run('\u6838\u5FC3\u673A\u5236\uFF1A').bold = True
doc.add_paragraph('useEffect \u5185\u8BBE\u7F6E setTimeout\uFF0C\u5EF6\u8FDF\u6267\u884C')
doc.add_paragraph('cleanup \u51FD\u6570\u8FD4\u56DE clearTimeout\uFF0C\u6E05\u9664\u65E7\u5B9A\u65F6\u5668')
doc.add_paragraph('\u6BCF\u6B21 value \u53D8\u5316\u2192\u5148\u6E05\u9664\u4E0A\u4E00\u4E2A\u5B9A\u65F6\u5668\u2192\u8BBE\u7F6E\u65B0\u5B9A\u65F6\u5668')
doc.add_paragraph('300ms \u5185\u5B89\u9759\u624D\u771F\u6B63\u66F4\u65B0\u503C')

doc.add_page_break()

# =================== Phase 5.3 ===================
doc.add_heading('\u4E09\u3001\u9636\u6BB5 5.3 \u2014 Zustand \u72B6\u6001\u7BA1\u7406', level=1)

doc.add_heading('3.1 \u4E3A\u4EC0\u4E48\u9700\u8981 Zustand', level=2)
p = doc.add_paragraph()
p.add_run('Prop Drilling\uFF1A').bold = True
p.add_run('\u72B6\u6001\u5728\u7EC4\u4EF6\u6811\u4E2D\u9010\u5C42\u4F20\u9012\uFF0C\u4E2D\u95F4\u5C42\u4E0D\u7528\u7684\u4E5F\u8981\u8F6C\u624B\u3002')

p = doc.add_paragraph()
p.add_run('Context \u7684\u95EE\u9898\uFF1A').bold = True
p.add_run('\u4E00\u4E2A Context \u7684\u503C\u53D8\u5316\u2192\u6240\u6709\u4F7F\u7528\u8BE5 Context \u7684\u7EC4\u4EF6\u90FD\u91CD\u6E32\u67D3\uFF0C\u5373\u4F7F\u4F60\u53EA\u7528\u5230\u5176\u4E2D\u4E00\u4E2A\u5B57\u6BB5\u3002')

p = doc.add_paragraph()
p.add_run('Zustand \u7684\u89E3\u51B3\u65B9\u6848\uFF1A').bold = True
p.add_run('\u9009\u62E9\u6027\u8BA2\u9605\uFF0C\u53EA\u8BA2\u9605\u4F60\u9700\u8981\u7684\u5B57\u6BB5\uFF0C\u53EA\u6709\u4F60\u8BA2\u9605\u7684\u5B57\u6BB5\u53D8\u5316\u65F6\u624D\u91CD\u6E32\u67D3\u3002')

doc.add_heading('3.2 Zustand \u6838\u5FC3\u6982\u5FF5', level=2)
doc.add_paragraph('Zustand \u662F\u7B2C\u4E09\u65B9\u5E93\uFF081KB\uFF0C\u96F6\u4F9D\u8D56\uFF09')
doc.add_paragraph('create() \u521B\u5EFA store\uFF0C\u8FD4\u56DE\u4E00\u4E2A hook \u51FD\u6570')
doc.add_paragraph('useStore() \u8FD4\u56DE\u4E00\u4E2A\u5BF9\u8C61\uFF0C\u6309\u540D\u5B57\u89E3\u6784\u53D6\u503C')

doc.add_heading('3.3 set \u7684\u4E24\u79CD\u5199\u6CD5', level=2)
p = doc.add_paragraph()
p.add_run('\u76F4\u63A5\u4F20\u503C\uFF1A').bold = True
p.add_run('\u65B0\u503C\u4E0D\u4F9D\u8D56\u65E7\u503C')
doc.add_paragraph("set({ isPlayingAll: false })")
p = doc.add_paragraph()
p.add_run('\u51FD\u6570\u5F62\u5F0F\uFF1A').bold = True
p.add_run('\u65B0\u503C\u4F9D\u8D56\u65E7\u503C\uFF08\u53D6\u6700\u65B0\u7684\u72B6\u6001\u5FEB\u7167\uFF09')
doc.add_paragraph("set((state) => ({ playAllIndex: state.playAllIndex + 1 }))")

doc.add_heading('3.4 pnpm \u5305\u7BA1\u7406\u5668', level=2)
doc.add_paragraph('pnpm \u662F npm \u7684\u66F4\u5FEB\u66FF\u4EE3\u54C1\uFF0C\u5B89\u88C5\u4F4D\u7F6E\u5728\u9879\u76EE\u5185\u7684 node_modules/')
doc.add_paragraph('node_modules/ \u2248 .venv/ | package.json \u2248 pyproject.toml | pnpm-lock.yaml \u2248 uv.lock')

doc.add_page_break()

# =================== Hands-on ===================
doc.add_heading('\u56DB\u3001\u5B9E\u6218 \u2014 \u8FC1\u79FB play-all \u5230 Zustand', level=1)

doc.add_heading('4.1 \u65B0\u5EFA Store', level=2)
doc.add_paragraph('\u6587\u4EF6\uFF1Alib/stores/play-all-store.ts')
doc.add_paragraph('\u5B9A\u4E49 playAllState \u63A5\u53E3\uFF0C\u542B isPlayingAll\u3001playAllIndex \u4E24\u4E2A\u72B6\u6001\u548C\u56DB\u4E2A\u64CD\u4F5C')
doc.add_paragraph('resetPlayAll() \u4E00\u6B65\u91CD\u7F6E\u4E24\u4E2A\u72B6\u6001\uFF0C\u4EE3\u66FF\u539F\u6765\u4E24\u884C\u4EE3\u7801')
doc.add_paragraph('nextInPlayAll() \u7528\u51FD\u6570\u5F62\u5F0F\u7684 set \u53D6\u6700\u65B0\u72B6\u6001\u6267\u884C +1')

doc.add_heading('4.2 \u6539\u9020 page.tsx', level=2)
p = doc.add_paragraph()
p.add_run('\u5220\u9664\uFF1A').bold = True
p.add_run('useState\u3001useCallback \u7684 memoized \u51FD\u6570')
p = doc.add_paragraph()
p.add_run('\u66FF\u6362\u4E3A\uFF1A').bold = True
p.add_run('const { isPlayingAll, ... } = usePlayAllStore()')
doc.add_paragraph('\u6240\u6709 setIsPlayingAll(false); setPlayAllIndex(0); \u2192 resetPlayAll()')
p = doc.add_paragraph()
p.add_run('\u6CE8\u610F\uFF1A').bold = True
p.add_run('export default \u5BFC\u5165\u65F6\u4E0D\u52A0\u82B1\u62EC\u53F7 {}')

doc.add_heading('4.3 \u6539\u9020\u5B50\u7EC4\u4EF6', level=2)
doc.add_paragraph('youtube-player.tsx \u548C highlights-panel.tsx\uFF1A')
doc.add_paragraph('\u5220\u9664 props \u7C7B\u578B\u4E2D\u7684 isPlayingAll\u3001playAllIndex \u7B49\u56DB\u4E2A\u5C5E\u6027')
doc.add_paragraph('\u7EC4\u4EF6\u5185\u90E8\u76F4\u63A5\u8C03\u7528 usePlayAllStore()')
p = doc.add_paragraph()
p.add_run('\u53BB\u6389 ?.\uFF08\u53EF\u9009\u94FE\uFF09\uFF1A').bold = True
p.add_run('store \u91CC\u7684\u51FD\u6570\u4E00\u5B9A\u5B58\u5728\uFF0C\u4E0D\u5B58\u5728 undefined\uFF0C\u4E0D\u9700\u8981\u4FDD\u62A4')

doc.add_heading('4.4 \u77E5\u8BC6\u70B9\u603B\u7ED3', level=2)
doc.add_paragraph('\u6570\u7EC4\u89E3\u6784 vs \u5BF9\u8C61\u89E3\u6784\uFF1A\u6570\u7EC4\u770B\u4F4D\u7F6E []\uFF0C\u5BF9\u8C61\u770B\u540D\u5B57 {}')
doc.add_paragraph('\u53EF\u9009\u94FE ?.\uFF1AsetIsPlayingAll?.(false) = \u51FD\u6570\u4E0D\u5B58\u5728\u5C31\u4E0D\u8C03\u7528\uFF0C\u9632\u6B62\u62A5\u9519')
doc.add_paragraph('export default import\uFF1A\u4E0D\u52A0\u82B1\u62EC\u53F7\uFF0C\u540D\u5B57\u81EA\u5B9A\u4E49')
doc.add_paragraph('TypeScript \u7F16\u8BD1\u68C0\u67E5\uFF1Anpx tsc --noEmit \u53EA\u68C0\u67E5\u4E0D\u8F93\u51FA')

doc.add_page_break()

# =================== Phase 5.4 ===================
doc.add_heading('\u4E94\u3001\u9636\u6BB5 5.4 \u2014 Server Components vs Client Components', level=1)

doc.add_heading('5.1 \u6838\u5FC3\u533A\u522B', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['', 'Server Component', 'Client Component']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

data = [
    ['\u6807\u8BB0', '\u4E0D\u52A0 "use client"', '\u52A0 "use client"'],
    ['\u8FD0\u884C\u73AF\u5883', '\u670D\u52A1\u7AEF (Node.js)', '\u6D4F\u89C8\u5668'],
    ['\u53EF\u7528\u529F\u80FD', '\u65E0 Hook\u3001\u65E0\u4E8B\u4EF6\u3001\u65E0API', '\u5168\u90E8\u53EF\u7528'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

p = doc.add_paragraph()
p.add_run('\u539F\u5219\uFF1A').bold = True
p.add_run('\u80FD\u7528 Server Component \u5C31\u4E0D\u7528 Client Component\u3002')

doc.add_heading('5.2 \u7EC4\u4EF6\u62C6\u5206\u7B56\u7565', level=2)
doc.add_paragraph('\u628A\u7EAF\u5C55\u793A\u7684\u90E8\u5206\u62C6\u4E3A Server Component\uFF0C\u53EA\u6709\u9700\u8981\u4EA4\u4E92\u7684\u90E8\u5206\u52A0 "use client"\u3002')
doc.add_paragraph('Server Component \u53EF\u4EE5 import Client Component\uFF0C\u53CD\u8FC7\u6765\u4E0D\u884C\u3002')

doc.add_page_break()

# =================== Phase 5.5 ===================
doc.add_heading('\u516D\u3001\u9636\u6BB5 5.5 \u2014 \u52A8\u6001\u8DEF\u7531', level=1)

doc.add_heading('6.1 [videoId] \u7684\u542B\u4E49', level=2)
doc.add_paragraph('app/analyze/[videoId]/page.tsx \u4E2D\u7684 [videoId] \u662F\u52A8\u6001\u8DEF\u7531\u53C2\u6570')
doc.add_paragraph('/analyze/abc123 \u2192 videoId = "abc123"')
doc.add_paragraph('/analyze/xyz789 \u2192 videoId = "xyz789"')
doc.add_paragraph('\u4E00\u4E2A\u6587\u4EF6\u5904\u7406\u6240\u6709\u89C6\u9891\u9875\u9762\uFF0C\u901A\u8FC7 useParams() \u83B7\u53D6\u53C2\u6570')

doc.add_heading('6.2 \u4E0E Flask \u5BF9\u6BD4', level=2)
doc.add_paragraph('Flask: @app.route(\'/analyze/<video_id>\')')
doc.add_paragraph('Next.js: app/analyze/[videoId]/page.tsx')
doc.add_paragraph('\u8BED\u6CD5\u4E0D\u540C\uFF0C\u529F\u80FD\u4E00\u6837\u2014\u2014\u90FD\u662F\u201C\u4EFB\u610F\u503C\u586B\u8FD9\u4E2A\u4F4D\u7F6E\u201D\u3002')

doc.add_page_break()

# =================== Phase 5.6 ===================
doc.add_heading('\u4E03\u3001\u9636\u6BB5 5.6 \u2014 API Routes', level=1)

doc.add_heading('7.1 page.tsx vs route.ts', level=2)
doc.add_paragraph('app/.../page.tsx \u2192 \u8FD4\u56DE HTML \u9875\u9762')
doc.add_paragraph('app/api/.../route.ts \u2192 \u8FD4\u56DE JSON \u6570\u636E')
p = doc.add_paragraph()
p.add_run('\u5BF9\u6BD4 Flask\uFF1A').bold = True
doc.add_paragraph('page.tsx \u2248 return render_template()')
doc.add_paragraph('route.ts \u2248 return jsonify(data)')

doc.add_heading('7.2 GET/POST \u7EA6\u5B9A', level=2)
p = doc.add_paragraph()
p.add_run('Next.js \u7528\u5BFC\u51FA\u51FD\u6570\u540D\u533A\u5206\u8BF7\u6C42\u65B9\u6CD5\uFF1A')
doc.add_paragraph('export async function GET(request) { ... }')
doc.add_paragraph('export async function POST(request) { ... }')
p = doc.add_paragraph()
p.add_run('POST\u3001GET \u53EA\u662F\u53D8\u91CF\u540D\uFF08JS \u5141\u8BB8\u5927\u5199\u53D8\u91CF\uFF09\uFF0C\u662F Next.js \u7684\u7EA6\u5B9A\u2014\u2014\u6846\u67B6\u6309\u51FD\u6570\u540D\u5206\u53D1\u7ED9\u5BF9\u5E94\u8BF7\u6C42\u65B9\u6CD5\u3002')

doc.add_heading('7.3 withSecurity \u4E2D\u95F4\u4EF6', level=2)
p = doc.add_paragraph()
p.add_run('\u4E2D\u95F4\u4EF6 = \u5B89\u68C0\u95E8\uFF1A').bold = True
doc.add_paragraph('\u524D\u7AEF\u8BF7\u6C42 \u2192 withSecurity(handler, PRESET) \u2192 \u5B89\u5168\u68C0\u67E5\u2192 handler \u6267\u884C \u2192 \u8FD4\u56DE\u7ED3\u679C')
p = doc.add_paragraph()
p.add_run('\u7B2C\u4E00\u4E2A\u53C2\u6570\uFF1A').bold = True
p.add_run('\u4E1A\u52A1 handler\uFF08\u4F60\u7684\u903B\u8F91\uFF09')
p = doc.add_paragraph()
p.add_run('\u7B2C\u4E8C\u4E2A\u53C2\u6570\uFF1A').bold = True
p.add_run('\u9884\u8BBE\u5B89\u5168\u7B49\u7EA7\uFF0C\u51B3\u5B9A\u68C0\u67E5\u7684\u4E25\u683C\u7A0B\u5EA6')
doc.add_paragraph('AUTHENTICATED\uFF1A\u5168\u5957\u68C0\u67E5 + \u5FC5\u987B\u767B\u5F55')
doc.add_paragraph('PUBLIC\uFF1A\u57FA\u7840\u68C0\u67E5\uFF0C\u4E0D\u9700\u8981\u767B\u5F55')
doc.add_paragraph('\u516D\u9053\u68C0\u67E5\uFF08\u65B9\u6CD5\u2192\u767B\u5F55\u2192\u9650\u6D41\u2192\u4F53\u79EF\u2192CSRF\u2192\u5B89\u5168\u5934\uFF09\u90FD\u5728 withSecurity \u5185\u90E8\u81EA\u52A8\u5B8C\u6210\u3002')

doc.add_page_break()

# =================== Q&A ===================
doc.add_heading('\u9644\u5F55\uFF1A\u7528\u6237\u63D0\u95EE\u8BB0\u5F55', level=1)

qa_items = [
    ('\u95EE\uFF1AusePlayAllStore \u4E3A\u4EC0\u4E48\u80FD\u76F4\u63A5 const \u8D4B\u503C\uFF1F',
     '\u7B54\uFF1Acreate() \u5185\u90E8\u5B58\u4E86\u72B6\u6001\u5BF9\u8C61\uFF0CusePlayAllStore() \u8FD4\u56DE\u7684\u5C31\u662F\u90A3\u4E2A\u5BF9\u8C61\uFF0C\u548C useState(false) \u8FD4\u56DE [\u503C,\u51FD\u6570] \u4E00\u6837\u3002'),
    ('\u95EE\uFF1A\u8FD4\u56DE\u503C\u662F\u6309\u5199\u5165\u987A\u5E8F\u5417\uFF1F',
     '\u7B54\uFF1A\u4E0D\u662F\uFF0C\u5BF9\u8C61\u89E3\u6784\u6309\u540D\u5B57\u53D6\u503C\uFF0C\u548C\u987A\u5E8F\u65E0\u5173\u3002\u6570\u7EC4\u6309\u4F4D\u7F6E\uFF0C\u5BF9\u8C61\u6309\u540D\u5B57\u3002'),
    ('\u95EE\uFF1A\u53EF\u9009\u94FE ?. \u662F\u4EC0\u4E48\uFF1F',
     '\u7B54\uFF1AsetIsPlayingAll?.(false) = \u5982\u679C\u51FD\u6570\u4E0D\u5B58\u5728\u5C31\u4E0D\u8C03\u7528\u3002\u4ECE store \u53D6\u7684\u51FD\u6570\u4E00\u5B9A\u5B58\u5728\uFF0C\u6240\u4EE5\u53EF\u4EE5\u53BB\u6389 ?.\u3002'),
    ('\u95EE\uFF1APOST \u4E3A\u4EC0\u4E48\u53EF\u4EE5\u5F53\u51FD\u6570\u540D\uFF1F',
     '\u7B54\uFF1AJS \u91CC POST \u53EA\u662F\u5927\u5199\u53D8\u91CF\u540D\uFF0C\u5408\u6CD5\u3002\u53EB\u8FD9\u4E2A\u540D\u5B57\u662F Next.js \u7EA6\u5B9A\uFF0C\u6846\u67B6\u6309\u51FD\u6570\u540D\u5206\u53D1\u8BF7\u6C42\u3002'),
    ('\u95EE\uFF1AwithSecurity \u7684\u4E24\u4E2A\u53C2\u6570\u662F\u4EC0\u4E48\uFF1F',
     '\u7B54\uFF1A\u7B2C\u4E00\u4E2A\u662F\u4E1A\u52A1 handler\uFF0C\u7B2C\u4E8C\u4E2A\u662F\u9884\u8BBE\u5B89\u5168\u7B49\u7EA7\uFF0C\u51B3\u5B9A\u68C0\u67E5\u7684\u4E25\u683C\u7A0B\u5EA6\u3002'),
]

for q, a in qa_items:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    doc.add_paragraph(a)
    doc.add_paragraph()  # spacer

# Save
output_path = r'D:\Develop\longcut\docs\learning-notes-05-26.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
